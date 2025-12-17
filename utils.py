import streamlit as st
import base64
import io
from docx import Document
from docx.shared import Pt

def load_css():
    """
    Injects the 'STRATOS Premium' Dark & Gold CSS theme.
    """
    st.markdown("""
        <style>
        /* Main Background */
        .stApp {
            background-color: #0E1117;
            color: #FAFAFA;
        }
        
        /* Sidebar */
        [data-testid="stSidebar"] {
            background-color: #161B22;
            border-right: 1px solid #30363D;
        }
        
        /* Headers */
        h1, h2, h3 {
            font-family: 'Helvetica Neue', sans-serif;
            font-weight: 700;
            color: #FFD700 !important; /* Gold */
            letter-spacing: -0.5px;
        }
        
        /* Buttons */
        .stButton > button {
            background-color: #1F6FEB;
            color: white;
            border: none;
            border-radius: 8px;
            padding: 0.5rem 1rem;
            font-weight: 600;
            transition: all 0.3s ease;
        }
        .stButton > button:hover {
            background-color: #FFD700;
            color: black;
            transform: translateY(-2px);
            box-shadow: 0 4px 12px rgba(255, 215, 0, 0.3);
        }
        
        /* Inputs */
        .stTextInput > div > div > input {
            background-color: #0D1117;
            color: white;
            border: 1px solid #30363D;
            border-radius: 6px;
        }
        .stTextInput > div > div > input:focus {
            border-color: #FFD700;
            box-shadow: 0 0 0 1px #FFD700;
        }
        
        /* Cards / Containers */
        div[data-testid="stExpander"] {
            background-color: #161B22;
            border: 1px solid #30363D;
            border-radius: 8px;
        }
        
        /* Success/Info Messages */
        .stAlert {
            background-color: #161B22;
            border: 1px solid #30363D;
            color: #FAFAFA;
        }
        
        /* Custom Footer */
        footer {visibility: hidden;}
        
        </style>
    """, unsafe_allow_html=True)

def header(title, subtitle=None):
    # Digital Clock (Robust Iframe Implementation)
    import streamlit.components.v1 as components
    
    # HTML/JS for the clock
    clock_html = """
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {
                margin: 0;
                padding: 0;
                background-color: transparent;
                color: #FFD700; /* Gold */
                font-family: 'Courier New', monospace;
                font-weight: bold;
                font-size: 16px;
                display: flex;
                justify-content: flex-end; /* Align right */
                align-items: center;
                height: 100vh;
            }
            #clock {
                padding-right: 10px;
            }
        </style>
    </head>
    <body>
        <div id="clock">Loading...</div>
        <script>
            function updateClock() {
                const now = new Date();
                const options = { weekday: 'long', year: 'numeric', month: 'long', day: 'numeric' };
                const dateString = now.toLocaleDateString(undefined, options);
                const timeString = now.toLocaleTimeString();
                document.getElementById('clock').innerHTML = dateString + ' | ' + timeString;
            }
            setInterval(updateClock, 1000);
            updateClock();
        </script>
    </body>
    </html>
    """
    
    # Inject iframe with transparent background
    # Height must be small to fit in header area
    components.html(clock_html, height=40, scrolling=False)

    st.markdown(f"# ⚡ {title}")
    if subtitle:
        st.markdown(f"### *{subtitle}*")
    st.markdown("---")

def create_docx(markdown_text):
    """
    Converts simple Markdown text to a Word Document binary stream.
    Handles Headers (#) and basic text.
    """
    doc = Document()
    
    # Set style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = markdown_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line.replace('- ', ''), style='List Bullet')
        else:
            doc.add_paragraph(line)
            
    # Save to memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def track_usage(platform_type):
    """
    Tracks usage stats in Session State.
    platform_type: 'LinkedIn' (3 hrs saved) or 'Twitter' (1 hr saved) or 'Blog' (4 hrs)
    """
    if 'hours_saved' not in st.session_state: st.session_state['hours_saved'] = 0.0
    if 'content_count' not in st.session_state: st.session_state['content_count'] = 0
    
    saved = 0
    if platform_type == 'LinkedIn': saved = 3.0
    elif platform_type == 'Twitter': saved = 1.0
    elif platform_type == 'Blog': saved = 4.0
    else: saved = 0.5
    
    st.session_state['hours_saved'] += saved
    st.session_state['content_count'] += 1

def display_impact_metrics():
    """
    Displays the 'Stratos Impact' section with centered title, realistic growing stats, and a big like button.
    """
    st.markdown("---")
    
    # Centered Title
    st.markdown("<h2 style='text-align: center; color: #FFD700;'>🚀 Stratos Impact</h2>", unsafe_allow_html=True)
    
    # Initialize Session State
    if 'hours_saved' not in st.session_state: st.session_state['hours_saved'] = 0.0
    if 'content_count' not in st.session_state: st.session_state['content_count'] = 0
    if 'likes' not in st.session_state: st.session_state['likes'] = 0
    
    # --- Global State Management (Cross-Session Persistence) ---
    # Use st.cache_resource to share data across all users/tabs
    @st.cache_resource
    class GlobalMetrics:
        def __init__(self):
            # Anchor: Dec 1, 2025 (Fresh Launch)
            # This ensures we start low (near 7 users) and grow
            self.anchor_time = 1764547200 
            self.manual_likes = 0 # Likes added by users clicking the button
            
        def get_stats(self):
            import time
            current_time = time.time()
            elapsed_hours = (current_time - self.anchor_time) / 3600
            
            # 1. Realistic User Count
            # Start with 7 users. Grow by ~1 user every 5 hours
            base_users = 7
            growth_rate = 5 
            simulated_users = base_users + int(elapsed_hours / growth_rate)
            
            # 2. Global Likes
            # Base (Simulated) + Manual (Real Clicks)
            base_likes = int(simulated_users * 4.2)
            total_likes = base_likes + self.manual_likes
            
            # 3. Hours Saved
            avg_hours_saved = 2.5
            total_hours = int(simulated_users * avg_hours_saved)
            
            # 4. Content Generated
            avg_content = 5
            total_content = int(simulated_users * avg_content)
            
            return simulated_users, total_likes, total_hours, total_content

        def add_like(self):
            self.manual_likes += 1
            
    # Instantiate Global State
    global_state = GlobalMetrics()
    
    # Get Current Global Stats
    g_users, g_likes, g_hours, g_content = global_state.get_stats()
    
    # Add Current Session Contributions to Display
    # (So the user sees their own impact immediately added to the global total)
    # Note: global_state.manual_likes already includes confirmed likes from other sessions
    # We just need to ensure we don't double count if we just clicked (handled by rerun)
    
    # Layout: Metrics on Left (3 cols), Like Button on Right (1 col)
    c1, c2, c3, c4 = st.columns([1, 1, 1, 1.2])
    
    with c1:
        st.metric(label="Total Users", value=f"{g_users:,}", delta="Growing")
    with c2:
        st.metric(label="Total Likes", value=f"{g_likes:,}", delta="Community")
    with c3:
        # Add session hours to global hours for display
        display_hours = g_hours + int(st.session_state['hours_saved'])
        st.metric(label="Hours Saved", value=f"{display_hours:,} hrs", delta="Efficiency")
        
    with c4:
        st.markdown("<br>", unsafe_allow_html=True) # Spacing
        # Big Like Button Logic
        if st.button("❤️ Like Stratos", use_container_width=True):
            if st.session_state['likes'] < 3:
                st.session_state['likes'] += 1
                global_state.add_like() # Increment Global Counter
                st.toast("Thanks for the love! ❤️")
                st.rerun() # Rerun to show updated global count
            else:
                st.toast("Max likes reached for this session! Share the love instead! 🚀")
        
        # Centered Like Count
        st.markdown(f"<div style='text-align: center; color: gray; font-size: 0.8em;'>Session Likes: {st.session_state['likes']}/3</div>", unsafe_allow_html=True)

    # Your Personal Stats (Subtle)
    if st.session_state['hours_saved'] > 0:
        st.markdown(f"<p style='text-align: center; color: #00C851;'>You've saved <b>{st.session_state['hours_saved']} hours</b> this session!</p>", unsafe_allow_html=True)

    st.markdown("### 📢 Share Your Strategy")
    st.info("💡 **Pro Tip:** The best way to share is to copy your generated content directly!")
    
    # Social Share Links
    share_text = f"I just saved {st.session_state['hours_saved']} hours on my content strategy using Stratos AI! ⚡ #StratosAI #ContentStrategy"
    share_url = "https://stratos-app.streamlit.app"
    
    twitter_url = f"https://twitter.com/intent/tweet?text={share_text}&url={share_url}"
    linkedin_url = f"https://www.linkedin.com/sharing/share-offsite/?url={share_url}"
    
    cols = st.columns(2)
    with cols[0]:
        st.link_button("🐦 Tweet Your Savings", twitter_url, use_container_width=True)
    with cols[1]:
        st.link_button("💼 Share on LinkedIn", linkedin_url, use_container_width=True)
    
    # Last Used Signature
    from datetime import datetime
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    st.markdown(f"<div style='text-align: center; color: #505050; font-size: 0.7em; margin-top: 20px;'>Last Active: {now_str}</div>", unsafe_allow_html=True)
