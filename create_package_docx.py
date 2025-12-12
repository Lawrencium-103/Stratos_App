import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx_from_md(md_path, output_path):
    """
    Converts the Stratos Replication Package Markdown to a formatted DOCX.
    """
    if not os.path.exists(md_path):
        print(f"Error: File not found at {md_path}")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Title Style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Headers
        if line.startswith('# '):
            h = doc.add_heading(line.replace('# ', ''), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        
        # List Items
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        
        # Code Blocks (Simple handling)
        elif line.startswith('```'):
            continue # Skip the fence markers
        
        # Blockquotes (Sales Scripts)
        elif line.startswith('> '):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Pt(20)
            runner = p.runs[0]
            runner.italic = True
            runner.font.color.rgb = RGBColor(100, 100, 100) # Grey
            
        # Normal Text
        else:
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f"Successfully created: {output_path}")

if __name__ == "__main__":
    md_file = r"c:/Users/USER/.gemini/antigravity/brain/e1e03750-8302-4bf4-8ba5-85af28db2313/stratos_replication_package.md"
    docx_file = r"c:/Users/USER/.gemini/antigravity/brain/e1e03750-8302-4bf4-8ba5-85af28db2313/Stratos_Replication_Package.docx"
    create_docx_from_md(md_file, docx_file)
