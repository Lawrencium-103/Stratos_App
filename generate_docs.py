import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_docx_from_md(md_path, docx_path):
    # Read Markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # Title
    title = doc.add_heading('STRATOS AI: Master Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line.replace('* ', '').replace('- ', ''), style='List Bullet')
        elif line.startswith('1. '):
            doc.add_paragraph(line.split('. ', 1)[1], style='List Number')
        elif line.startswith('```'):
            continue # Skip code block markers for simplicity in this quick script
        else:
            doc.add_paragraph(line)
            
    doc.save(docx_path)
    print(f"Successfully created {docx_path}")

if __name__ == "__main__":
    md_file = r"c:/Users/USER/.gemini/antigravity/brain/e1e03750-8302-4bf4-8ba5-85af28db2313/stratos_master_documentation.md"
    docx_file = r"c:/Users/USER/.gemini/antigravity/brain/e1e03750-8302-4bf4-8ba5-85af28db2313/Stratos_Master_Documentation.docx"
    create_docx_from_md(md_file, docx_file)
